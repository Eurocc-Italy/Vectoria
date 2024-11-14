#
# VECTORIA
#
# @authors : Andrea Proia, Chiara Malizia, Leonardo Baroncelli
#
import re
import logging
from pathlib import Path
import docx
from langchain.docstore.document import Document
from vectoria_lib.common.config import Config
# Do not remove the follwing imports (they are used by globals())
from re import search, sub, match, fullmatch

logger = logging.getLogger('db_management')
config = Config()

def _apply_regex(text: str, regex: str, regex_function: str) -> str:
    compiled_regex = re.compile(regex)
    result = globals()[regex_function](compiled_regex, text) # TODO: not very generic
    if result:
        return result.group(0).strip() # TODO: not very generic
    raise ValueError(f"No match found for regex {regex} in text")

def _extract_metadata_from_unstructured_data(unstructured_data: list[Document], regexes: list[dict]) -> dict:
    if regexes is None: regexes = []
    raw_text = "".join([doc.page_content for doc in unstructured_data])
    metadata = {}
    for regex in regexes:
        metadata[regex["metadata_name"]] = _apply_regex(raw_text, regex["regex_pattern"], regex["regex_function"])
    return metadata

def extract_text_from_docx_file(
        file_path: Path,
        filter_paragraphs: list = [], # TODO: implement this feature
        dump_doc_structure_on_file: bool = False,
        regexes_for_metadata_extraction: list[dict] = []
) -> list[Document]:
    """
    Extracts text from a DOCX file. 
    The extraction process relies on the tags: "heading", "paragraph", and "table".
    Anything that is not a heading, paragraph, or table is considered unstructured data.

    You can pass a unstrucured data parser to the function, which will be used to parse the unstructured data.
    The unstructured data parser will extract metadata (as a dictionary) from the unstructured data. 
    The metadata will be added to the Document objects.

    In addition to unstructured data metadata, the function adds the following metadata keys:
    * doc_file_name: the name of the DOCX file
    * paragraph_number: the number of the paragraph
    * layout_tag: the tag of the layout element (Heading, Paragraph, Table)

    Returns a list of Document objects, containing page content and metadata.
    """
    logger.debug("Extracting text from %s", file_path.stem)

    document = docx.Document(file_path)

    document_flat_structure = _extract_flat_structure(document)

    paragraphs_numbers_and_names = _recover_paragraphs_numbers_and_names(document_flat_structure)

    document_flat_structure = _to_document_objects(document_flat_structure)

    document_flat_structure, paragraphs_numbers_and_names, unstructured_data = _filter_unstructured_data(document_flat_structure, paragraphs_numbers_and_names)

    metadata_from_unstructured_data = _extract_metadata_from_unstructured_data(
        unstructured_data,
        regexes_for_metadata_extraction
    )
    
    _add_metadata(
        document_flat_structure,
        paragraphs_numbers_and_names,
        metadata_from_unstructured_data,
        file_path.name
    )

    logger.debug("Extracted %d documents from %s", len(document_flat_structure), file_path.stem)

    if dump_doc_structure_on_file:
        _log_document_structure_on_file(document_flat_structure, Path(config.get("vectoria_logs_dir")) / "docs_structure", file_path.stem) 


    document_flat_structure = _filter_headings(document_flat_structure)

    document_flat_structure = _merge_paragraphs_contents(document_flat_structure)

    return document_flat_structure

def _merge_paragraphs_contents(document_flat_structure: list[Document]):
    merged_paragraphs = []
    
    paragraphs = sorted(list(set([doc.metadata["paragraph_number"] for doc in document_flat_structure])))

    for p in paragraphs:
        documents_with_same_paragraph_number = [doc for doc in document_flat_structure if doc.metadata["paragraph_number"] == p]

        merged_content = "\n".join([doc.page_content for doc in documents_with_same_paragraph_number])

        merged_paragraphs.append(Document(page_content=merged_content, metadata=documents_with_same_paragraph_number[0].metadata))

    return merged_paragraphs

def _filter_headings(document_flat_structure: list[Document]):
    return [doc for doc in document_flat_structure if "Heading" not in doc.metadata["layout_tag"]]    

def _extract_flat_structure(doc: str) -> list[tuple]:
    """
    Extracts the flat structure of the document.
    """
    current_heading = None  # Track the current heading
    structure = []  # Store content structure

    for element in doc.element.body:
        #print(element.tag, element.text)
        
        if element.tag.endswith('p'):  # It's a paragraph
            paragraph = docx.text.paragraph.Paragraph(element, doc)
    
            heading_level = _get_heading_level(paragraph)
            
            if heading_level:
                current_heading = paragraph.text  # Update current heading
                structure.append((f"Heading {heading_level}", paragraph.text))
            else:
                if paragraph.text:
                    structure.append(("Paragraph", paragraph.text))
        
        elif element.tag.endswith('tbl'):  # It's a table
            table_data = _extract_table_data(element, doc)
            if not _check_table_empty(table_data):
                structure.append(("Table", table_data))

    return structure

def _recover_paragraphs_numbers_and_names(flat_structure: list[tuple]) -> list[str]:
    heading_levels = []

    # To store the result
    result = []

    element_text = ""
    for element in flat_structure:
        element_type = element[0]

        if element_type.startswith("Heading"):
            element_text = element[1]

            # Extract the level of heading (e.g., 'Heading 2' -> level 2)
            level = int(element_type.split()[1])

            # Ensure the heading_levels list has the correct size
            while len(heading_levels) < level:
                heading_levels.append(0)
            
            # Increment the correct level, reset lower levels if needed
            heading_levels[level - 1] += 1
            heading_levels = heading_levels[:level]

            # Generate the heading number string
            heading_number = ".".join(map(str, heading_levels))
            result.append((heading_number, element_text))
        
        elif element_type == 'Paragraph' or element_type == 'Table':
            # Paragraphs or tables inherit the current heading numbering
            current_number = ".".join(map(str, heading_levels))
            result.append((current_number, element_text))

    #breakpoint()
    return result

def _to_document_objects(document_flat_structure: list[tuple]) -> list[Document]:
    return [
        Document(
            page_content=str(element[1]), 
            metadata={"layout_tag": element[0]}
        ) 
        for element in document_flat_structure
    ]

def _filter_unstructured_data(docs: list[Document], paragraphs_numbers_and_names: list[tuple[str, str]]):
    """
    Everything Document objects that do not have a paragraph number are considered unstructured data
    and they are removed from the list.
    The unstructured data is returned as a Document object.
    """
    unstructured_data = []
    docs_to_keep = []
    paragraphs_numbers_to_keep = []
    for doc, number_and_name in zip(docs, paragraphs_numbers_and_names):
        if number_and_name[0] == "":
            unstructured_data.append(Document(page_content=str(doc.page_content), metadata={}))
        else:
            docs_to_keep.append(doc)
            paragraphs_numbers_to_keep.append(number_and_name)
    return docs_to_keep, paragraphs_numbers_to_keep, unstructured_data

def _add_metadata(
        document_flat_structure: list[Document],
        paragraphs_numbers_and_names: list[tuple[str, str]],
        metadata_from_unstructured_data: dict,
        doc_file_name: str
) -> None:
    if len(document_flat_structure) != len(paragraphs_numbers_and_names):
        raise ValueError(f"The number of documents {len(document_flat_structure)} and the number of paragraphs numbers {len(paragraphs_numbers_and_names)} do not match")
    
    for doc, number_and_name in zip(document_flat_structure, paragraphs_numbers_and_names):
        doc.metadata["doc_file_name"] = doc_file_name
        doc.metadata["paragraph_number"] = number_and_name[0]
        doc.metadata["paragraph_name"] = number_and_name[1]
        doc.metadata.update(metadata_from_unstructured_data)

def _log_document_structure_on_file(
        docs: list[Document], 
        log_in_folder: Path, 
        output_file_prefix: str
) -> None:
    Path(log_in_folder).mkdir(parents=True, exist_ok=True)
    file_path = Path(log_in_folder) / f"{output_file_prefix}_structure.txt"
    with open(file_path, "w", encoding="utf-8") as f:
        # Print the document structure, including where tables are located
        for doc in docs:
            if doc.metadata["layout_tag"].startswith("Heading"):
                print(f"{doc.metadata['paragraph_number']}   {doc.metadata['layout_tag']}: {doc.page_content}", file=f)
            elif doc.metadata["layout_tag"] == "Paragraph":
                print(f"{doc.metadata['paragraph_number']}   Paragraph: {doc.page_content}", file=f)
            elif doc.metadata["layout_tag"] == "Table":
                print(f"{doc.metadata['paragraph_number']}   Table under Heading: {doc.page_content}", file=f)

def _get_heading_level(paragraph): 
    # A helper function to determine the level of a heading based on style
    if paragraph.style.name.startswith("Heading"):
        return int(paragraph.style.name.split()[-1])  # Extract heading level number
    return None  # Not a heading

def _check_table_empty(table):
    # Check if a table is empty
    for row in table:
        for col in row:
            if col:
                return False
    return True

# def _extract_table_data(table, doc):
#     table = docx.table.Table(table, doc)
#     table_data = []
#     for row in table.rows:
#         row_data = [cell.text for cell in row.cells]
#         table_data.append(row_data)
#     return table_data

# TODO: verificare che l'LLM capisca la tabella in questo formato (senza "[ [], [] ]")
def _extract_table_data(table, doc):
    table = docx.table.Table(table, doc)
    table_data = ""
    for row in table.rows:
        row_data = "".join([f"{cell.text} " for cell in row.cells])
        table_data += "\n" + row_data
    return table_data