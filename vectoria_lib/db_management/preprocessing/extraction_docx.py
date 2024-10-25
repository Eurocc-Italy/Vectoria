#
# VECTORIA
#
# @authors : Andrea Proia, Chiara Malizia, Leonardo Baroncelli
#

import logging
from pathlib import Path
from typing import Callable
import docx
from langchain.docstore.document import Document

from vectoria_lib.db_management.preprocessing.document_data import  DocumentData

logger = logging.getLogger('db_management')
    
def extract_text_from_docx_file(
        file_path: Path,
        filter_paragraphs: list,
        log_in_folder: Path = None,
        unstructured_data_parser: Callable = None
) -> list[Document]:
    """
    Extracts text from a DOCX file. 
    The extraction process relies on the tags: "heading", "paragraph", and "table".
    Anything that is not a heading, paragraph, or table is considered unstructured data.

    You can pass a unstrucured data parser to the function, which will be used to parse the unstructured data.
    The unstructured data parser will extract metadata from the unstructured data. 
    
    The metadata will be added to the Document objects.
    Returns a list of Document objects, containing page content and metadata.
    """
    logger.debug("Extracting text from %s", file_path.stem)

    print(file_path)

    metadata = {
        "source": file_path.name
    }
    document = docx.Document(file_path)

    document_flat_structure = _extract_flat_structure(document)
    paragraphs_numbers      = _recover_paragraphs_numbers(document_flat_structure)
    document_flat_structure = _to_document_objects(document_flat_structure)

    unstructured_data = _filter_unstructured_data(document_flat_structure, paragraphs_numbers)
    if unstructured_data_parser is not None:
        logger.debug("Parsing unstructured data")
        metadata.update(unstructured_data_parser(unstructured_data))
    
    _add_metadata(document_flat_structure, paragraphs_numbers, metadata)

    logger.debug("Extracted %d documents from %s", len(document_flat_structure), file_path.stem)

    _log_document_structure_on_file(document_flat_structure, log_in_folder, file_path.stem) 
    
    return document_flat_structure

def _to_document_objects(document_flat_structure: list[tuple]) -> list[Document]:
    return [
        Document(
            page_content=str(element[1]), 
            metadata={"type": element[0]}
        ) 
        for element in document_flat_structure
    ]

def _add_metadata(
        document_flat_structure: list[Document],
        paragraphs_numbers: list[str],
        metadata: dict
) -> None:
    for doc, number in zip(document_flat_structure, paragraphs_numbers):
        doc.metadata["source"] = metadata["source"]
        doc.metadata["number"] = number

def _filter_unstructured_data(docs: list[Document], paragraphs_numbers: list[str]) -> Document:
    """
    Everything Document objects that do not have a paragraph number are considered unstructured data
    and they are removed from the list.
    The unstructured data is returned as a Document object.
    """
    c = 0
    unstructured_data = Document(page_content="", metadata={"type": "unstructured"})
    for doc, number in zip(docs, paragraphs_numbers):
        if number == "":
            unstructured_data.page_content += str(doc.page_content)
            c += 1
            docs.remove(doc)
    logger.debug("Found %d unstructured data", c)

    return unstructured_data

def _log_document_structure_on_file(docs: list[Document], log_in_folder: Path, output_file_prefix: str):
    if log_in_folder is not None:
        Path(log_in_folder).mkdir(parents=True, exist_ok=True)
        file_path = Path(log_in_folder) / f"{output_file_prefix}_structure.txt"
        with open(file_path, "w", encoding="utf-8") as f:
            # Print the document structure, including where tables are located
            for doc in docs:
                if doc.metadata["type"].startswith("Heading"):
                    print(f"{doc.metadata['number']}   {doc.metadata['type']}: {doc.page_content}", file=f)
                elif doc.metadata["type"] == "Paragraph":
                    print(f"{doc.metadata['number']}   Paragraph: {doc.page_content}", file=f)
                elif doc.metadata["type"] == "Table":
                    print(f"{doc.metadata['number']}   Table under Heading: {doc.page_content}", file=f)

def _extract_flat_structure(doc: str) -> list[tuple]:
    """
    Extracts the flat structure of the document.
    """
    current_heading = None  # Track the current heading
    structure = []  # Store content structure

    for element in doc.element.body:
        #print(element.tag, element.text)
        
        if element.tag.endswith('p'):  # It's a paragraph
            paragraph = docx.text.paragraph.Paragraph(element, doc)
    
            heading_level = _get_heading_level(paragraph)
            
            if heading_level:
                current_heading = paragraph.text  # Update current heading
                structure.append((f"Heading {heading_level}", paragraph.text))
            else:
                if paragraph.text:
                    structure.append(("Paragraph", paragraph.text))
        
        elif element.tag.endswith('tbl'):  # It's a table
            table_data = _extract_table_data(element, doc)
            if not _check_table_empty(table_data):
                structure.append(("Table", table_data, current_heading))

    return structure

def _recover_paragraphs_numbers(flat_structure: list[tuple]) -> list[str]:
    heading_levels = []

    # To store the result
    result = []

    for element in flat_structure:
        element_type = element[0]

        if element_type.startswith("Heading"):
            # Extract the level of heading (e.g., 'Heading 2' -> level 2)
            level = int(element_type.split()[1])

            # Ensure the heading_levels list has the correct size
            while len(heading_levels) < level:
                heading_levels.append(0)
            
            # Increment the correct level, reset lower levels if needed
            heading_levels[level - 1] += 1
            heading_levels = heading_levels[:level]

            # Generate the heading number string
            heading_number = ".".join(map(str, heading_levels))
            result.append(heading_number)
        
        elif element_type == 'Paragraph' or element_type == 'Table':
            # Paragraphs or tables inherit the current heading numbering
            current_number = ".".join(map(str, heading_levels))
            result.append(current_number)
        
        else:
            # For any other types like 'Table', process as needed
            result.append(".".join(map(str, heading_levels)))

    return result

def _get_heading_level(paragraph): 
    # A helper function to determine the level of a heading based on style
    if paragraph.style.name.startswith("Heading"):
        return int(paragraph.style.name.split()[-1])  # Extract heading level number
    return None  # Not a heading

def _check_table_empty(table):
    # Check if a table is empty
    for row in table:
        for col in row:
            if col:
                return False
    return True

def _extract_table_data(table, doc):
    table = docx.table.Table(table, doc)
    table_data = []
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        table_data.append(row_data)
    return table_data
    
def _extract_text_structure(document_flat_structure: list[tuple], paragraphs_numbers: list[tuple]):
    """
    ######## From:

    document_flat_structure = [
        ('Paragraph', 'Questa è benzina'),
        ('Paragraph', 'non strutturata.'),
        ('Heading 1', 'INTRODUZIONE'), # capitolo 1
        ('Heading 2', 'Scopo'), # capitolo 1.1
        ('Paragraph', 'Lo scopo del presente documento è la descrizione delle regole per la gestione, l’elaborazione ed emissione dei documenti del sistema normativo della Divisione Elettronica.'),
        ('Paragraph', 'Il BMS è lo strumento per la pubblicazione, la diffusione e l’archiviazione dei documenti di origine interna, i modelli/moduli ad essi associati, la documentazione emessa da Corporate e ogni altro documento esterno con particolare riferimento a leggi internazionali e standard, che si ritenga opportuno diffondere, purché non coperti da copyright di terze parti.'),
        ('Heading 2', 'Applicabilità'), # capitolo 1.2
        ('Paragraph', 'Il presente documento si applica alla Divisione Elettronica (perimetro Italia).'),
        ('Heading 1', 'RIFERIMENTI'), # capitolo 2
        ('Heading 2', 'Documenti'), # capitolo 2.1
        ('Paragraph', 'I documenti sotto riportati sono da intendere nell’ultima versione correntemente emessa, se non diversamente riportato.'),
        ('Table', [['Codice', 'Titolo'], ['AQAP-2110', 'NATO Quality Assurance Requirement for Design Development and Production'], ['AQAP-2210', 'NATO Supplementary Software Quality Assurance Requirement to AQAP2110 or AQAP2310']], 'Documenti'),
        ('Paragraph', 'NOTA: I documenti indicati con (*) sono in fase di emissione alla data di pubblicazione del presente documento. Nel transitorio si applicano i corrispondenti documenti Legacy.')
    ]

    ######### to:
    
    [{'childs': [{'childs': [],
                'name': 'Scopo',
                'docs': [Document('Lo scopo del presente documento è la descrizione delle '
                        'regole per la gestione, l’elaborazione ed emissione '
                        'dei documenti del sistema normativo della Divisione '
                        'Elettronica.'),
                        Document('Il BMS è lo strumento per la pubblicazione, la '
                        'diffusione e l’archiviazione dei documenti di origine '
                        'interna, i modelli/moduli ad essi associati, la '
                        'documentazione emessa da Corporate e ogni altro '
                        'documento esterno con particolare riferimento a leggi '
                        'internazionali e standard, che si ritenga opportuno '
                        'diffondere, purché non coperti da copyright di terze '
                        'parti.')]},
                {'childs': [],
                'name': 'Applicabilità',
                'docs': [Document('Il presente documento si applica alla Divisione '
                        'Elettronica (perimetro Italia).')]}],
    'name': 'INTRODUZIONE',
    'docs': []},
    {'childs': [{'childs': [],
                'name': 'Documenti',
                'docs': [Document('I documenti sotto riportati sono da intendere '
                        'nell’ultima versione correntemente emessa, se non '
                        'diversamente riportato.'),
                        Document("[['Codice', 'Titolo'], ['AQAP-2110', 'NATO Quality "
                        'Assurance Requirement for Design Development and '
                        "Production'], ['AQAP-2210', 'NATO Supplementary "
                        'Software Quality Assurance Requirement to AQAP2110 or '
                        "AQAP2310']]"),
                        Document('NOTA: I documenti indicati con (*) sono in fase di '
                        'emissione alla data di pubblicazione del presente '
                        'documento. Nel transitorio si applicano i '
                        'corrispondenti documenti Legacy.')]}],
    'name': 'RIFERIMENTI',
    'docs': []}]
    """

        
    structured_data = []
    unstructured_data = Document(page_content="", metadata={"type": "unstructured"})
    heading_stack = []  # This stack will keep track of current headings at different levels

    # Find for unstructured data first
    for item in document_flat_structure:
        if not item[0].startswith('Heading'):
            unstructured_data.page_content += str(item[1])
        else:
            break
        
    current_name = None
    current_level = None
    current_id = 0
    
    for paragraph_number, item in zip(paragraphs_numbers, document_flat_structure):

        # Check if the item is a heading (e.g., 'Heading 1', 'Heading 2', etc.)
        if item[0].startswith('Heading'):
            level = int(item[0].split()[1])  # Extract heading level (e.g., 1, 2, 3, etc.)
            current_name = item[1]
            current_level = level
            current_id = 0
            heading_dict = {
                "name": item[1],
                "number": paragraph_number,
                "doc": None,
                "childs": []
            }
            
            # Adjust the stack based on heading level
            while len(heading_stack) >= level:

                heading_stack.pop()  # Go up to the correct parent level by popping

            if heading_stack:
                
                # Add the new heading as a child of the current top of the stack
                heading_stack[-1]["childs"].append(heading_dict)
            else:
                # If the stack is empty, it means we're at the top level
                structured_data.append(heading_dict)
                
            # Push the new heading onto the stack
            heading_stack.append(heading_dict)
        
        # If the item is a paragraph or table, add it to the current section
        elif item[0] == 'Paragraph' or item[0] == 'Table':
            text = str(item[1])
            if heading_stack:
                if heading_stack[-1]["doc"] is None:
                    heading_stack[-1]["doc"] = Document(page_content=text, metadata={"name": current_name, "level": current_level, "id": current_id})
                    current_id += 1
                else:
                    heading_stack[-1]["doc"].page_content += text
                            

    return structured_data, unstructured_data

def _get_flat_docs_list(structured_data) -> list[Document]:
    # recursively extract all the Document objects from the children of structured_data
    docs = []
    for item in structured_data:
        if item["doc"] is not None:
            docs.append(
                item["doc"]
            )
        docs.extend(_get_flat_docs_list(item["childs"]))
    return docs