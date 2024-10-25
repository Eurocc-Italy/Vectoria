import re
import pytest
from pathlib import Path
from langchain.docstore.document import Document
import docx
from vectoria_lib.db_management.preprocessing.extraction_docx import (
    extract_text_from_docx_file,
    _extract_flat_structure,
    _recover_paragraphs_numbers,
    _filter_unstructured_data,
    _to_document_objects
)
from vectoria_lib.common.paths import TEST_DIR
from vectoria_lib.common.config import Config

def test_extract_flat_structure_from_googledocs():
    """
    Google Docs docx file does not support the tbl tag.
    """
    document = docx.Document(TEST_DIR / "data/docx/docx_from_googledocs.docx")
    flat_structure = _extract_flat_structure(document)
    assert len(flat_structure) == 20
    assert flat_structure == [
        ('Paragraph', 'Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data\n\nSome unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data Some unstructured data'), 
        ('Paragraph', 'Test document 2 title'), 
        ('Paragraph', 'Next a summary:'), 
        ('Heading 1', 'First chapter (Heading 1)'), 
        ('Paragraph', 'First Lorem ipsum 🥰'), 
        ('Paragraph', 'Now spaces and empty lines:'), 
        ('Paragraph', '     '), 
        ('Paragraph', 'End first chapter.'), 
        ('Heading 1', 'Second Chapter (Heading 1)'), 
        ('Paragraph', 'Second Lorem ipsum '), 
        ('Heading 2', 'Sub-chapter (Heading 2)'), 
        ('Paragraph', 'Now: line and formula and table:'), 
        ('Paragraph', 'End second chapter'), 
        ('Heading 1', 'Third Chapter'), 
        ('Paragraph', 'Here’s an image'), 
        ('Heading 3', 'Image'), 
        ('Heading 1', 'Final hierarchy (Heading 1)'), 
        ('Heading 2', 'Heading 2'), 
        ('Heading 3', 'Heading 3'), 
        ('Heading 4', 'Heading 4')
    ]

def test_extract_flat_structure_from_word():
    """
    Word docx file supports the tbl tag.
    """
    document = docx.Document(TEST_DIR / "data/docx/docx_from_word.docx")
    flat_structure = _extract_flat_structure(document)
    assert flat_structure == [
        ('Paragraph', 'Document Title'),
        ('Paragraph', 'Here’s a summary:'),
        ('Heading 1', 'Title 1'),
        ('Paragraph', 'Content of title 1: Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.'), 
        ('Heading 2', 'Title 2'),
        ('Paragraph', 'Content of title 2: Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.'), 
        ('Heading 3', 'Title 3'),
        ('Paragraph', 'Content of title 3: Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.'), 
        ('Heading 4', 'Title 4'),
        ('Paragraph', 'Content of title 4: Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.'), 
        ('Heading 5', 'Title 5'),
        ('Paragraph', 'Content of title 5: pefffforza!'),
        ('Heading 1', 'Another title 1'),
        ('Paragraph', 'Here’s a table:'),
        ('Table', [['Row1 Col1', 'Row1 Col2'], ['Row2 Col1', 'Row2 Col2']], 'Another title 1'), 
        ('Heading 2', 'Another title 2'),
        ('Paragraph', 'Here’s an image:')
    ]

def test_recover_paragraphs_numbers():
    document = docx.Document(TEST_DIR / "data/docx/docx_from_word.docx")
    flat_structure = _extract_flat_structure(document)
    paragraphs_numbers = _recover_paragraphs_numbers(flat_structure)
    assert paragraphs_numbers == [
        '', '', '1', '1', '1.1', '1.1', '1.1.1', '1.1.1', '1.1.1.1', '1.1.1.1', '1.1.1.1.1', '1.1.1.1.1', '2', '2', '2', '2.1', '2.1'
    ]
    
def test_filter_unstructured_data():
    document = docx.Document(TEST_DIR / "data/docx/docx_from_word.docx")
    flat_structure = _extract_flat_structure(document)
    paragraphs_numbers = _recover_paragraphs_numbers(flat_structure)
    docs = _to_document_objects(flat_structure)
    docs_to_keep, paragraphs_numbers_to_keep, unstructured_data = _filter_unstructured_data(docs, paragraphs_numbers)
    assert len(docs_to_keep) == 15
    assert len(paragraphs_numbers_to_keep) == 15
    assert isinstance(unstructured_data, Document)

def test_extract_text_from_docx_file():
    config = Config()
    config.load_config(TEST_DIR / "data" / "config" / "test_config.yaml")

    docs: list[Document] = extract_text_from_docx_file(
        TEST_DIR / "data/docx/docx_from_word.docx",
        filter_paragraphs=[],
        dump_doc_structure_on_file=True,
        regexes_for_metadata_extraction = [{
            "metadata_name": "first_symbol",
            "regex_pattern": r"^[A-Za-z]",
            "regex_function": "search"
        }]
    )

    assert isinstance(docs, list)
    assert isinstance(docs[0], Document)
    assert len(docs) == 15

    assert Path(config.get("vectoria_logs_dir") / "docs_structure" / "docx_from_word_structure.txt").exists()

    assert set(docs[0].metadata.keys()) == set(["layout_tag", "paragraph_number", "doc_file_name", "first_symbol"])
    assert docs[0].metadata["first_symbol"] == "D"

    """
    id = re.search(r"IDENTIFICATIVO\s*:\s*(.*)", text).group(1).strip()

    date = re.search(r"DATA\s*:\s*(.*)", text).group(1).strip()
    
    doc_type = re.search(r"TIPO DOCUMENTO\s*:\s*(.*)", text).group(1).strip()
    
    app = re.search(r"APPLICAZIONE\s*:\s*(.*)", text).group(1).strip()
    summary = re.search(r"SOMMARIO\s*:\s*(.*(?:\n.*)*)", text).group(1).strip()
    """