from pathlib import Path
import docx
from langchain.docstore.document import Document
from vectoria_lib.ingestion.extraction_docx import (
    extract_text_from_docx_file,
    _extract_flat_structure,
    _recover_paragraphs_numbers_and_names,
    _filter_unstructured_data,
    _to_document_objects
)
from vectoria_lib.common.paths import TEST_DIR
from vectoria_lib.common.config import Config

def test_extract_flat_structure_from_word(config):
    """
    Word docx file supports the tbl tag.
    """
    document = docx.Document(TEST_DIR / "data/docx/docx_extraction_test.docx")
    flat_structure = _extract_flat_structure(document)

    assert flat_structure == [
        ('Paragraph', 'Document Title'),
        ('Paragraph', 'Here’s a summary:'),
        ('Heading 1', 'Title 1'),
        ('Paragraph', 'Content of title 1: Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.'), 
        ('Heading 2', 'Title 2'),
        ('Paragraph', 'Content of title 2: Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.'), 
        ('Heading 3', 'Title 3'),
        ('Paragraph', 'Content of title 3: Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.'), 
        ('Heading 4', 'Title 4'),
        ('Paragraph', 'Content of title 4: Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.'), 
        ('Heading 5', 'Title 5'),
        ('Paragraph', 'Content of title 5: pefffforza!'),
        ('Heading 1', 'Another title 1'),
        ('Paragraph', 'Here’s a table:'),
        ('Table', '\nRow1 Col1 Row1 Col2 \nRow2 Col1 Row2 Col2 '),
        ('Heading 2', 'Another title 2'),
        ('Paragraph', 'Here’s an image:')
    ]

def test_recover_paragraphs_numbers_and_names(config):
    document = docx.Document(TEST_DIR / "data/docx/docx_extraction_test.docx")
    flat_structure = _extract_flat_structure(document)
    paragraphs_numbers_and_names = _recover_paragraphs_numbers_and_names(flat_structure)

    assert paragraphs_numbers_and_names == [
        ('', ''), ('', ''), ('1', 'Title 1'), ('1', 'Title 1'), ('1.1', 'Title 2'), ('1.1', 'Title 2'), ('1.1.1', 'Title 3'), ('1.1.1', 'Title 3'), ('1.1.1.1', 'Title 4'), ('1.1.1.1', 'Title 4'), ('1.1.1.1.1', 'Title 5'), ('1.1.1.1.1', 'Title 5'), ('2', 'Another title 1'), ('2', 'Another title 1'), ('2', 'Another title 1'), ('2.1', 'Another title 2'), ('2.1', 'Another title 2')
    ]
    
def test_filter_unstructured_data(config):
    document = docx.Document(TEST_DIR / "data/docx/docx_extraction_test.docx")
    flat_structure = _extract_flat_structure(document)
    paragraphs_numbers = _recover_paragraphs_numbers_and_names(flat_structure)
    docs = _to_document_objects(flat_structure)
    docs_to_keep, paragraphs_numbers_to_keep, unstructured_data = _filter_unstructured_data(docs, paragraphs_numbers)
    assert len(docs_to_keep) == 15
    assert len(paragraphs_numbers_to_keep) == 15
    assert len(unstructured_data) == 2
    assert isinstance(unstructured_data[0], Document)

def test_extract_text_from_docx_file(config):
    docs: list[Document] = extract_text_from_docx_file(
        TEST_DIR / "data/docx/docx_extraction_test.docx",
        dump_doc_structure_on_file=True,
        regexes_for_metadata_extraction = [{
            "name": "first_symbol",
            "pattern": r"^[A-Za-z]"
        }]
    )   

    assert isinstance(docs, list)
    assert isinstance(docs[0], Document)
    assert len(docs) == 7

    assert Path( Path(config.get("vectoria_logs_dir")) / "docs_structure" / "docx_extraction_test_structure.txt").exists()

    assert set(docs[0].metadata.keys()) == set(["layout_tag","paragraph_name", "paragraph_number", "doc_file_name", "first_symbol"])
    assert docs[0].metadata["first_symbol"] == "D"
    for d in docs:
        assert "Heading" not in d.metadata["layout_tag"]